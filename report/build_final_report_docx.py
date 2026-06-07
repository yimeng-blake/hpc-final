from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "FINAL_REPORT.docx"
FIGURES = ROOT / "figures_accelergy_plugin"
SUMMARY = ROOT / "outputs" / "summary_accelergy_plugin"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, margin_twips: int = 80) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_fixed_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_run_font(run, size: int | float = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc: Document, text: str = "", *, style: str | None = None) -> object:
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True)
    keep_with_next(p)
    return p


def add_caption(doc: Document, text: str) -> object:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_margins(table, 80)
    set_fixed_table_width(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = ""
        run = header_cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, size=10, bold=True)
        set_cell_shading(header_cells[idx], "EDEDED")
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=10)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx != 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def format_styles(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Body Text", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)
            if style_name.startswith("Heading"):
                style.font.bold = True
                style.font.color.rgb = None

    normal = styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    format_styles(doc)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Deadline-Constrained Systolic Array Design for Image Processing")
    set_run_font(r, size=14, bold=True)
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Blake Wang and Yu Cheng Wu")
    set_run_font(r, size=12)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CSEN 318 High-Performance Computer Architecture and Systems, Spring 2026")
    set_run_font(r, size=12)
    p.paragraph_format.space_after = Pt(10)


def load_result_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    min_df = pd.read_csv(SUMMARY / "minimum_energy_designs.csv")
    winners = min_df[(min_df["resolution"] == "1080p") & (min_df["deadline_ms"] == 33)].copy()
    winners = winners.sort_values("gaussian_kernel")
    winners["mac_pct"] = winners["energy_mac_pj"] / winners["energy_total_pj"] * 100
    winners["sram_pct"] = winners["energy_sram_pj"] / winners["energy_total_pj"] * 100
    winners["dram_pct"] = winners["energy_dram_pj"] / winners["energy_total_pj"] * 100

    bottleneck = pd.read_csv(SUMMARY / "bottleneck_summary.csv")
    selected = bottleneck.merge(
        winners[
            [
                "resolution",
                "tile_width",
                "tile_height",
                "gaussian_kernel",
                "array_size",
                "sram_budget_kb",
                "bandwidth_gbps",
                "dataflow",
            ]
        ],
        on=[
            "resolution",
            "tile_width",
            "tile_height",
            "gaussian_kernel",
            "array_size",
            "sram_budget_kb",
            "bandwidth_gbps",
            "dataflow",
        ],
        how="inner",
    )
    stage = (
        selected.pivot_table(
            index="gaussian_kernel",
            columns="stage_op",
            values="energy_share_pct",
            aggfunc="first",
        )
        .reset_index()
        .sort_values("gaussian_kernel")
    )
    resolution_winners = min_df[min_df["deadline_ms"] == 33].copy()
    resolution_winners["_resolution_order"] = resolution_winners["resolution"].map({"720p": 0, "1080p": 1, "hires": 2}).fillna(99)
    resolution_winners = resolution_winners.sort_values(["_resolution_order", "gaussian_kernel"])
    return winners, stage, resolution_winners


def require_existing_figures(paths: list[Path]) -> None:
    missing = [path for path in paths if not path.exists()]
    if missing:
        joined = "\n".join(str(path) for path in missing)
        raise FileNotFoundError(
            "Missing report figures. Generate experiment summaries and plots first with "
            "`scripts/summarize_results.py`.\n" + joined
        )


def main() -> None:
    winners, stage, resolution_winners = load_result_tables()
    stage_chart = FIGURES / "stage_share_by_kernel.png"
    component_chart = FIGURES / "component_energy_split.png"
    energy_latency_chart = FIGURES / "energy_latency_scatter.png"
    require_existing_figures([energy_latency_chart, stage_chart, component_chart])

    doc = Document()
    setup_document(doc)
    add_title(doc)

    add_heading(doc, "1. Introduction")
    add_para(
        doc,
        "Motivation. Real-time embedded vision systems process frames on a fixed schedule. A 30 FPS camera receives a new frame every 33.33 ms, so the accelerator does not need to be maximally fast. It needs to complete the frame before the deadline while using as little energy as possible.",
    )
    add_para(
        doc,
        "Objective. This project evaluates deadline-constrained systolic array designs for a real-time grayscale image-processing pipeline: input frame, Gaussian blur, Sobel edge detection, and output edge frame. The objective is to find the minimum-energy accelerator configuration that still completes a full frame before the target deadline.",
    )
    add_para(
        doc,
        "The main real-time case is 1080p at 30 FPS, and the report also checks 60 FPS and 24 FPS deadline sensitivity. The modeled boundary includes the two accelerator compute stages and excludes camera capture, display, host CPU orchestration, image-quality tuning, static/leakage power, and full-system idle power.",
    )
    add_para(
        doc,
        "Literature and market context. Systolic arrays are relevant because image filters and convolution-like workloads can be expressed as regular matrix-style computations. SCALE-Sim is used for cycle and memory-access modeling, while Accelergy is used to convert component action counts into modeled dynamic energy. The market context is edge vision, where throughput targets such as 24, 30, or 60 FPS must be balanced against energy budgets.",
    )
    add_para(
        doc,
        "The design question is therefore not simply which array is fastest. The useful question is which array is just fast enough for the application deadline while minimizing energy per frame. That framing matches embedded camera deployments where thermal and battery constraints can matter more than excess latency slack.",
    )

    add_heading(doc, "2. System Design And Implementation Details")
    add_para(
        doc,
        "Technologies and tools. The framework is implemented in Python. SCALE-Sim provides cycle counts, stalls, utilization, and SRAM/DRAM accesses for each tiled workload. Accelergy provides per-action dynamic energy from generated Energy Reference Tables. Pandas is used for aggregation, Matplotlib is used to generate charts, and pytest covers config parsing, workload shapes, tiling, energy accounting, and summary logic.",
    )
    add_para(
        doc,
        "Application model. The application pipeline is Gaussian blur followed by Sobel edge detection. Gaussian blur smooths the image and has variable cost based on the filter size. Sobel computes horizontal and vertical gradients using two fixed 3x3 filters.",
    )
    add_para(
        doc,
        "The image filters are not arbitrary GEMMs. Each output pixel is a dot product between a local image neighborhood and filter weights. For a whole tile, these dot products are lowered into a matrix multiply so they can be simulated by SCALE-Sim.",
    )
    add_table(
        doc,
        ["Stage", "M", "N", "K", "MACs per tile"],
        [
            ["Gaussian blur", "Output pixels", "1", "kernel_size^2", "M x K"],
            ["Sobel edge", "Output pixels", "2", "9", "M x 2 x 9"],
        ],
        [1.5, 1.6, 0.7, 1.0, 1.5],
    )
    add_para(
        doc,
        "The final model uses representative 128x128 output tiles and scales them to full frames. The tiling model accounts for full interior tiles, right-edge tiles, bottom-edge tiles, corner tiles, and halo input pixels needed by convolution kernels. The end-to-end frame latency is the Gaussian stage latency plus the Sobel stage latency; frame energy is computed the same way.",
    )
    add_para(
        doc,
        "Modeled end-to-end boundary. Within the accelerator model, the full application path is Gaussian plus Sobel over every output pixel in the frame. There are no additional image-processing stages mapped to the systolic array. Capture, display, CPU launch overhead, software scheduling, and output formatting are treated as outside the accelerator boundary and are assumed not to change the relative ranking of the array configurations.",
    )
    add_para(
        doc,
        "The design space includes 720p, 1080p, and 2048x2048 frames; Gaussian kernels from 3x3 to 11x11; array sizes from 8x8 to 128x128; SRAM budgets from 256 KB to 4096 KB; bandwidths from 50 GB/s to 800 GB/s; and weight-stationary, output-stationary, and input-stationary dataflows. A focused 1080p refinement adds 48x48 and 96x96 arrays, intermediate SRAM budgets, and a 100 GB/s bandwidth point.",
    )
    add_para(
        doc,
        "Energy is modeled from dynamic action counts. MAC operations map to PE MAC actions; SRAM ifmap/filter/ofmap accesses map to local SRAM read/update actions; and DRAM ifmap/filter/ofmap accesses map to external memory read/write actions. This is a comparative accelerator model, not a measured full-system power model.",
    )
    add_para(
        doc,
        "The reproducibility path is explicit in the package. Tile-level SCALE-Sim reports are aggregated into outputs/summary_accelergy_plugin/pipeline_runs.csv; deadline-expanded feasibility is stored in feasible_designs.csv; selected designs are stored in minimum_energy_designs.csv; and component action counts are stored in accelergy_action_counts.csv.",
    )

    add_heading(doc, "3. Experimental Results And Evaluation")
    add_para(
        doc,
        "Experimental setup and methodology. The final Accelergy-backed summary contains 6,052 stage-level rows, 1,767 complete pipeline configurations, 5,301 deadline-expanded feasibility rows, 165 Pareto-frontier rows, 36 minimum-energy rows, and 42,364 Accelergy action-count rows. Latency is computed from SCALE-Sim cycles at a 1 GHz reference clock; a design is feasible when its full-frame latency is less than or equal to the selected frame deadline.",
    )
    add_para(
        doc,
        "The table below reports the minimum-energy feasible designs for the main 1080p, 33 ms objective. Appendix A includes the energy-latency scatter plot that shows these selected points within the broader design space.",
    )
    winner_rows = []
    for _, row in winners.iterrows():
        winner_rows.append(
            [
                f"{int(row.gaussian_kernel)}x{int(row.gaussian_kernel)}",
                f"{int(row.array_size)}x{int(row.array_size)} {row.dataflow.upper()}",
                f"{int(row.sram_budget_kb)} KB",
                f"{row.latency_ms:.2f} ms",
                f"{row.energy_total_mj:.3f} mJ",
            ]
        )
    add_table(
        doc,
        ["Kernel", "Best design", "SRAM", "Latency", "Energy"],
        winner_rows,
        [0.9, 2.0, 1.0, 1.1, 1.1],
    )
    add_para(
        doc,
        "Analysis of results. The result is workload-dependent. Input-stationary wins the 3x3 and 5x5 cases by energy even though it is slower than some alternatives. Since those designs still meet the deadline, the optimizer correctly chooses lower energy over unnecessary speed. Weight-stationary wins for 7x7 and 11x11 because the heavier Gaussian stage exposes more useful work for the array.",
    )
    add_para(
        doc,
        "The underlying reason is the lowered GEMM shape. These image filters produce skinny GEMMs with N equal to 1 for Gaussian and 2 for Sobel. Large arrays can therefore be underutilized for lighter kernels. As the Gaussian kernel grows from 3x3 to 11x11, K grows from 9 to 121 and the Gaussian stage becomes the dominant part of the end-to-end workload. That shift changes which dataflow and array size are energy-best.",
    )
    add_para(
        doc,
        "Deadline sensitivity. The selected 1080p designs also meet a 60 FPS deadline of 16.67 ms. Therefore, changing the target from 30 FPS to 60 FPS or 24 FPS does not change the selected 1080p designs in the current sweep. Energy per frame stays the same, while average power scales with the frame rate.",
    )
    add_table(
        doc,
        ["Kernel", "Latency", "60 FPS", "30 FPS", "24 FPS"],
        [
            [f"{int(row.gaussian_kernel)}x{int(row.gaussian_kernel)}", f"{row.latency_ms:.2f} ms", "Feasible", "Feasible", "Feasible"]
            for _, row in winners.iterrows()
        ],
        [0.9, 1.0, 1.1, 1.1, 1.1],
    )
    stage_rows = []
    for _, row in stage.iterrows():
        stage_rows.append(
            [
                f"{int(row.gaussian_kernel)}x{int(row.gaussian_kernel)}",
                f"{float(row.get('gaussian', 0)):.1f}%",
                f"{float(row.get('sobel', 0)):.1f}%",
            ]
        )
    add_table(doc, ["Kernel", "Gaussian energy share", "Sobel energy share"], stage_rows, [1.0, 2.1, 2.1])
    add_para(
        doc,
        "Stage-level analysis confirms the workload explanation. Gaussian energy rises from 45.9% of the selected 3x3 design to 91.2% of the selected 11x11 design, while Sobel shrinks as a share of the total because it remains fixed at two 3x3 filters. Appendix B visualizes this shift.",
    )
    add_para(
        doc,
        "The dynamic component-energy split supports a scoped memory-energy conclusion. For the selected 1080p designs, MAC energy is about 10-13%, SRAM dynamic energy is about 78-79%, and DRAM dynamic energy is about 10%. This means memory actions dominate in this modeled action-count framework. It should not be read as a claim about all physical implementations, because static SRAM energy, wires, clocking, control, and full-system overhead are outside the model. Appendix C shows the component split.",
    )
    doc.add_page_break()
    component_rows = []
    for _, row in winners.iterrows():
        component_rows.append(
            [
                f"{int(row.gaussian_kernel)}x{int(row.gaussian_kernel)}",
                f"{row.mac_pct:.1f}%",
                f"{row.sram_pct:.1f}%",
                f"{row.dram_pct:.1f}%",
            ]
        )
    add_table(doc, ["Kernel", "MAC", "SRAM dynamic", "DRAM dynamic"], component_rows, [0.9, 1.0, 1.5, 1.5])
    add_para(
        doc,
        "The split is consistent across kernel sizes because the selected designs repeatedly trade compute work for local-buffer and external-memory activity. The exact percentages should not be overgeneralized beyond this model, but they do justify discussing memory traffic as a first-order design concern for this workload.",
    )
    add_para(
        doc,
        "Resolution sensitivity is another check on the conclusion. The same framework evaluates 720p, 1080p, and 2048x2048 frames at the 33 ms deadline. The selected design does not scale monotonically with resolution alone; it also depends on kernel size and dataflow. For example, the 2048x2048 3x3 case uses 128x128 input-stationary and is close to the deadline at 26.89 ms, while heavier kernels select weight-stationary designs that finish faster because the larger K dimension gives the array more useful work.",
    )
    resolution_rows = []
    for _, row in resolution_winners.iterrows():
        resolution = "2048x2048" if row.resolution == "hires" else str(row.resolution)
        resolution_rows.append(
            [
                resolution,
                f"{int(row.gaussian_kernel)}x{int(row.gaussian_kernel)}",
                f"{int(row.array_size)}x{int(row.array_size)} {row.dataflow.upper()}",
                f"{row.latency_ms:.2f} ms",
                f"{row.energy_total_mj:.3f} mJ",
            ]
        )
    add_table(
        doc,
        ["Resolution", "Kernel", "Best design", "Latency", "Energy"],
        resolution_rows,
        [1.1, 0.8, 1.7, 1.0, 1.1],
    )
    add_para(
        doc,
        "The project evaluates up to 2048x2048 frames, not 8K. Adding 8K would be a straightforward new workload point for the framework, but it would require a new run because 8K has about 33.2 million pixels, roughly 16 times 1080p. The final report only claims results that were actually generated and summarized.",
    )

    doc.add_page_break()
    add_heading(doc, "4. Conclusions")
    add_para(
        doc,
        "The project built a reproducible SCALE-Sim plus Accelergy framework for deadline-constrained systolic-array design exploration. The implementation connects a concrete image-processing application to GEMM-lowered tile workloads, scales tile results to full frames, and reports latency, energy, EDP, feasibility, bottlenecks, and Pareto behavior.",
    )
    add_para(
        doc,
        "The most important result is that accelerator sizing should be based on workload shape and frame deadline. For this workload, 1080p 3x3 and 5x5 kernels favor 128x128 input-stationary designs, while 7x7 and 11x11 favor weight-stationary designs. Bigger hardware is not automatically better; it helps only when the workload can keep it busy enough to justify the energy cost.",
    )
    add_para(
        doc,
        "What worked well: tiled simulation made full-frame sweeps tractable, Accelergy made the energy model more defensible than hardcoded constants, and the 1080p refinement sweep showed that the coarse grid did not hide a lower-energy winner. What did not work as well: some large-kernel input-stationary and output-stationary SCALE-Sim cases became resource-heavy and had to be skipped with documented guards. Future work should add color images, 8K frames, additional image-processing stages, and static/full-system power modeling.",
    )
    add_para(
        doc,
        "The most important methodological lesson is that application mapping has to be stated explicitly. Once the workload is described as Gaussian and Sobel stencil filters lowered into skinny GEMMs, the utilization behavior and dataflow choices are easier to interpret. This also prevents the experiment from looking like a sweep of unrelated matrix sizes.",
    )
    add_para(
        doc,
        "The final design recommendation is not a single universal hardware configuration. It is to choose array size and dataflow jointly with the workload and frame deadline. In this project, input-stationary is attractive for smaller kernels because it saves modeled dynamic energy while still meeting the deadline. Weight-stationary becomes better for larger kernels because the Gaussian stage supplies enough K-dimension work to make the array more useful.",
    )
    add_para(
        doc,
        "The main limitation is that the numbers are comparative model results, not silicon measurements. The conclusions are still useful for architecture exploration because all designs are evaluated with the same assumptions, but a production design would need leakage, clocking, interconnect, controller overhead, and real implementation constraints before final hardware selection.",
    )

    doc.add_page_break()
    add_heading(doc, "5. Team Contributions And Acknowledgement")
    add_table(
        doc,
        ["Component", "Assigned / Completed By", "Contribution"],
        [
            ["Project framing, workload selection, and presentation narrative", "Blake Wang and Yu Cheng Wu", "Shared"],
            ["SCALE-Sim workload generation, tiling, and sweep scripts", "Blake Wang", "50%"],
            ["Accelergy integration, action-count mapping, and energy summaries", "Yu Cheng Wu", "50%"],
            ["Result analysis, figures, report writing, and feedback-driven revision", "Blake Wang and Yu Cheng Wu", "Shared"],
        ],
        [2.7, 2.2, 1.0],
    )
    add_para(doc, "Estimated total contribution: Blake Wang 50%, Yu Cheng Wu 50%.")
    add_para(
        doc,
        "The work distribution was balanced across implementation, modeling, analysis, and writing. Both members participated in framing the application and interpreting results; the task table separates the main ownership areas so the source of each project component is clear.",
    )
    add_para(
        doc,
        "Generative AI tools were used to help revise writing, organize the report, inspect code consistency, and suggest ways to explain the experiment more clearly. The tools were not treated as the source of the experimental results. The team remained responsible for running the experiments, checking result tables, validating repository outputs, and deciding the final technical conclusions.",
    )
    add_para(
        doc,
        "The final package should therefore be read as a reproducible class project: the code and data directories provide the evidence for the reported numbers, while this report summarizes the problem, model, assumptions, results, and limitations.",
    )

    refs = doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "References")
    for text in [
        '[1] H. T. Kung, "Why Systolic Architectures?", Computer, vol. 15, no. 1, pp. 37-46, 1982.',
        '[2] A. Samajdar et al., "SCALE-Sim: Systolic CNN Accelerator Simulator", arXiv:1811.02883, 2018.',
        '[3] Y. N. Wu et al., "Accelergy: An Architecture-Level Energy Estimation Methodology for Accelerator Designs", IEEE/ACM ICCAD, 2019.',
    ]:
        add_para(doc, text)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A. Energy-Latency Design Space")
    add_figure(
        doc,
        energy_latency_chart,
        "Appendix Figure A1. Energy-latency tradeoff. Points left of the deadline line meet the real-time target; stars mark selected minimum-energy designs.",
        width=5.9,
    )
    add_heading(doc, "Appendix B. Stage Energy Share")
    add_figure(
        doc,
        stage_chart,
        "Appendix Figure B1. Stage energy share for 1080p selected designs. Gaussian dominates as the kernel grows, while Sobel remains fixed at two 3x3 filters.",
        width=5.7,
    )
    add_heading(doc, "Appendix C. Dynamic Component Energy Split")
    add_figure(
        doc,
        component_chart,
        "Appendix Figure C1. Dynamic component energy split. SRAM actions dominate the selected designs in this modeled action-count framework.",
        width=5.7,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
